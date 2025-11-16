from typing import Dict
from pathlib import Path
import io

def parse_file(file_content: bytes, filename: str) -> Dict[str, any]:
    """
    Parse uploaded file and extract text content
    
    Args:
        file_content: Raw file bytes
        filename: Name of the uploaded file
        
    Returns:
        Dict with success status and parsed content or error
    """
    try:
        # Get file extension
        file_ext = Path(filename).suffix.lower()
        
        # Text files (.txt, .md)
        if file_ext in ['.txt', '.md']:
            try:
                text = file_content.decode('utf-8')
                return {
                    'success': True,
                    'content': text,
                    'file_type': 'text',
                    'filename': filename
                }
            except UnicodeDecodeError:
                return {
                    'success': False,
                    'error': f'Unable to decode {file_ext} file. Please ensure it is valid UTF-8 text.'
                }
        
        # PDF files
        elif file_ext == '.pdf':
            try:
                import PyPDF2
                
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                if len(pdf_reader.pages) == 0:
                    return {
                        'success': False,
                        'error': 'PDF file has no pages.'
                    }
                
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                        continue
                
                if not text.strip():
                    return {
                        'success': False,
                        'error': 'PDF appears to be empty or contains only images/scanned content. Please upload a text-based PDF.'
                    }
                
                return {
                    'success': True,
                    'content': text,
                    'file_type': 'pdf',
                    'filename': filename,
                    'pages': len(pdf_reader.pages)
                }
                
            except ImportError:
                return {
                    'success': False,
                    'error': 'PyPDF2 library not installed. Please install it to process PDF files.'
                }
            except Exception as e:
                return {
                    'success': False,
                    'error': f'Failed to parse PDF: {str(e)}. The file may be corrupted or password-protected.'
                }
        
        # DOCX files
        elif file_ext == '.docx':
            try:
                import docx
                
                doc = docx.Document(io.BytesIO(file_content))
                
                # Extract text from paragraphs
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                
                # Extract text from tables
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        table_text.append(" | ".join(row_text))
                
                text = "\n".join(paragraphs)
                if table_text:
                    text += "\n\nTables:\n" + "\n".join(table_text)
                
                if not text.strip():
                    return {
                        'success': False,
                        'error': 'DOCX file appears to be empty.'
                    }
                
                return {
                    'success': True,
                    'content': text,
                    'file_type': 'docx',
                    'filename': filename,
                    'paragraphs': len(paragraphs),
                    'tables': len(doc.tables)
                }
                
            except ImportError:
                return {
                    'success': False,
                    'error': 'python-docx library not installed. Please install it to process DOCX files.'
                }
            except Exception as e:
                return {
                    'success': False,
                    'error': f'Failed to parse DOCX: {str(e)}. The file may be corrupted.'
                }
        
        # Unsupported file type
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_ext}. Please upload .txt, .md, .pdf, or .docx files.'
            }
    
    except Exception as e:
        return {
            'success': False,
            'error': f'Unexpected error parsing file: {str(e)}'
        }


def analyze_content(content: str) -> Dict[str, any]:
    """
    Analyze parsed content for quality and suitability
    
    Args:
        content: Extracted text content
        
    Returns:
        Dict with analysis results and validation status
    """
    # Strip whitespace
    content_stripped = content.strip()
    
    # Check if empty
    if not content_stripped:
        return {
            'valid': False,
            'error': 'File content is empty.'
        }
    
    # Count words and characters
    words = content_stripped.split()
    word_count = len(words)
    char_count = len(content_stripped)
    
    # Minimum length check
    if char_count < 50:
        return {
            'valid': False,
            'error': f'Content too short ({char_count} characters). Please upload a file with at least 50 characters of study material.'
        }
    
    if word_count < 10:
        return {
            'valid': False,
            'error': f'Content too short ({word_count} words). Please upload a file with at least 10 words.'
        }
    
    # Maximum length check
    if char_count > 100000:
        return {
            'valid': False,
            'error': f'Content too long ({char_count} characters). Please upload a file under 100,000 characters to ensure optimal processing.'
        }
    
    # Check for reasonable text structure
    lines = [line.strip() for line in content_stripped.split('\n') if line.strip()]
    line_count = len(lines)
    
    # Estimate reading time (average 200 words per minute)
    reading_time_minutes = word_count / 200
    
    return {
        'valid': True,
        'word_count': word_count,
        'char_count': char_count,
        'line_count': line_count,
        'estimated_reading_time': round(reading_time_minutes, 1),
        'content_preview': content_stripped[:200] + ('...' if len(content_stripped) > 200 else '')
    }


def validate_audio_transcription(transcription: str) -> Dict[str, any]:
    """
    Validate transcribed audio for quality and completeness
    
    Args:
        transcription: Transcribed text from speech
        
    Returns:
        Dict with validation results
    """
    if not transcription or not transcription.strip():
        return {
            'valid': False,
            'error': 'No speech detected. Please try speaking again.'
        }
    
    transcription_stripped = transcription.strip()
    words = transcription_stripped.split()
    word_count = len(words)
    
    # Minimum length check
    if word_count < 3:
        return {
            'valid': False,
            'error': f'Response too short ({word_count} words). Please provide a more detailed explanation (at least 10 words).'
        }
    
    if word_count < 10:
        return {
            'valid': False,
            'error': f'Response is brief ({word_count} words). For a good explanation, try to use at least 10 words to explain the concept.',
            'suggestion': 'Try explaining it as if you were teaching it to someone who has never heard of it before.'
        }
    
    # Check for gibberish or non-sensical patterns
    # Very simple heuristic: check if words are too short on average
    avg_word_length = sum(len(word) for word in words) / word_count if word_count > 0 else 0
    
    if avg_word_length < 2:
        return {
            'valid': False,
            'error': 'Audio transcription may be unclear. Please try speaking more clearly.',
            'suggestion': 'Speak at a moderate pace in a quiet environment.'
        }
    
    return {
        'valid': True,
        'word_count': word_count,
        'transcription': transcription_stripped,
        'quality': 'good' if word_count >= 20 else 'acceptable'
    }
